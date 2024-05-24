import docx
import json

filename_de = 'static/data/docx/cbt_programm_de.docx'
filename_hu = 'static/data/docx/cbt_programm_hu.docx'
filename_sk = 'static/data/docx/cbt_programm_sk.docx'
filename_pl = 'static/data/docx/cbt_programm_pl.docx'
filename_cs = 'static/data/docx/cbt_programm_cs.docx'


def create_json(filename, lang='de'):
    doc = docx.Document(filename)
    paragraphs = doc.paragraphs

    json_filename = 'static/json/' + filename.split('/')[-1].split('.')[0] + '.json'
    print(json_filename)

    json_data = {}

    tag = ''
    zeit = ''
    titel = ''
    content = ''
    location = ''

    counter = 0

    # try:
    if (True):
        i = -1
        while i < len(paragraphs) - 1:
            i += 1
            print(i)
            text = paragraphs[i].text
            runs = paragraphs[i].runs

            if text.startswith('***'):
                break

            elif text.startswith('Freitag') or text.startswith('Pátek') or text.endswith('péntek') or text.startswith(
                    'Piątek') or text.startswith('Piatok'):
                tag = 0
                zeit = ''
            elif text.startswith('Samstag') or text.startswith('Sobota') or text.endswith('szombat') or text.startswith(
                    'Sobota') or text.startswith('Sobota'):
                tag = 1
                zeit = ''
            elif text.startswith('Sonntag') or text.startswith('Neděle') or text.endswith(
                    'vasárnap') or text.startswith('Niedziela') or text.startswith('nedeľa'):
                tag = 2
                zeit = ''

            elif len(text) > 0 and text[0].isdigit():
                zeit = text

            elif '\t' in text:
                print(text)
                if paragraphs[i + 1] and len(paragraphs[i + 1].runs) > 0 and paragraphs[i + 1].runs[0].bold and \
                        paragraphs[i + 1].runs[0].italic:
                    print(text, paragraphs[i + 1].text)
                    i += 1
                continue

            elif len(runs) > 0 and runs[0].bold and runs[-1].bold:
                if paragraphs[i - 1] and paragraphs[i + 2]:
                    if paragraphs[i - 1].text == '' and paragraphs[i + 2].text == '':
                        if paragraphs[i + 1].text != '':
                            location = {'name': text.strip(), 'address': paragraphs[i + 1].text.strip()}
                        continue

                if text.lower().startswith('program'):
                    continue
                if zeit == '':
                    continue

                titel = text
                if json_data.get(counter) and json_data[counter]['content-' + lang] == '':
                    json_data[counter]['titel-' + lang] += '\n' + titel
                else:
                    counter += 1
                    json_data[counter] = {}
                    json_data[counter]['tag'] = tag
                    json_data[counter]['zeit'] = zeit.strip()
                    json_data[counter]['location-' + lang] = location
                    json_data[counter]['titel-' + lang] = titel.strip()
                    json_data[counter]['content-' + lang] = ''

            elif len(runs) > 0 and not runs[0].bold and runs[0].italic and runs[-1].italic and json_data[counter].get(
                    'content-' + lang):
                print(text)
                if len(text) < 100:
                    location = {'name': text.strip(), 'address': ""}
                    json_data[counter]['location-' + lang] = location


            else:
                content = text.strip()
                if text == '':
                    continue

                if json_data.get(counter) and json_data[counter].get('content-' + lang):

                    next_paragraph = paragraphs[i + 1]

                    content_split = content.split(', ')
                    all_same_length = all(len(s) == 2 for s in content_split)

                    if (counter < 22
                            and len(content) > 3
                            and not all_same_length
                            and next_paragraph != None
                            and (next_paragraph.text == '' or len(next_paragraph.text) == 2)
                            and (len(paragraphs[i - 1].runs) == 0 or not paragraphs[i - 1].runs[0].bold)):
                        location = {'name': content, 'address': ''}
                        json_data[counter]['location-' + lang] = location
                        continue

                    if len(content) == 2 or all_same_length:
                        json_data[counter]['lang'] = content_split
                        continue

                    json_data[counter]['content-' + lang] += '\n' + content.strip()

                else:
                    json_data[counter]['content-' + lang] = content.strip()

    # except Exception as e:
    #    print(e)

    print()
    print(counter)

    with open(json_filename, 'w') as f:
        json.dump(json_data, f, indent=4, ensure_ascii=False)
        f.close()


def create_json2(filename, lang='de'):
    doc = docx.Document(filename)
    paragraphs = doc.paragraphs

    json_filename = 'static/json/' + filename.split('/')[-1].split('.')[0] + '.json'
    print(json_filename)

    json_data = {}

    tag = ''
    zeit = ''
    location = ''

    counter = 0

    def check_block(block_arg):
        nonlocal tag, zeit, location, counter
        titel = ''
        untertitel = ''
        content = ''

        program_lang = []

        if len(block_arg) == 0:
            return

        text = block_arg[0].text

        if text.startswith('Freitag') or text.startswith('Pátek') or text.endswith('péntek') or text.startswith(
                'Piątek') or text.startswith('Piatok'):
            tag = 0
            zeit = ''
            return check_block(block_arg[1:])
        elif text.startswith('Samstag') or text.startswith('Sobota') or text.endswith('szombat') or text.startswith(
                'Sobota') or text.startswith('Sobota'):
            tag = 1
            zeit = ''
            return check_block(block_arg[1:])
        elif text.startswith('Sonntag') or text.startswith('Neděle') or text.endswith(
                'vasárnap') or text.startswith('Niedziela') or text.startswith('nedeľa'):
            tag = 2
            zeit = ''
            return check_block(block_arg[1:])

        new_block = block_arg

        if block_arg[0].text[0].isdigit():
            zeit = text
            new_block = block_arg[1:]
        elif len(block_arg) > 1 and block_arg[1].text[0].isdigit():
            location = {'name': text.strip()}
            zeit = block_arg[1].text
            new_block = block_arg[2:]


        for pos in range(len(new_block)):

            text = new_block[pos].text
            runs = new_block[pos].runs

            if '\t' in text or '\\t' in text:
                return

            elif runs[0].bold:
                if text.lower().startswith('program'):
                    return

                if len(block_arg) == 2 and len(new_block) == 2 and len(runs) > 0 and not new_block[1].runs[0].bold:
                    location = {'name': new_block[pos + 1].text.strip(), 'bezeichnung': text.strip(), }
                    return

                if titel == '':
                    titel = text.strip()
                else:
                    if runs[0].italic:
                        if untertitel != '':
                            untertitel += '\n' + text.strip()
                        else:
                            untertitel = text.strip()
                    else:
                        titel += '\n' + text.strip()

            elif runs[0].italic and content != '':
                if 'englisch' in text.lower() or 'angol' in text.lower() or 'angličtine' in text.lower() or 'angielskim' in text.lower() or 'angličtině' in text.lower():
                    if content != '':
                        content += '\n' + text.strip()
                    else:
                        content = text.strip()
                    continue
                if pos >= len(new_block) - 2:
                    location = {'name': text.strip()}
                else:
                    if content != '':
                        content += '\n' + text.strip()
                    else:
                        content = text.strip()

            else:
                if content != '':

                    content_split = text.split(', ')
                    all_same_length = all(len(s) == 2 for s in content_split)

                    if all_same_length:
                        program_lang = content_split
                        continue

                    content += '\n' + text.strip()

                else:
                    content = text.strip()

        if titel == '':
            return

        counter += 1
        json_data[counter] = {}
        json_data[counter]['tag'] = tag
        json_data[counter]['zeit'] = zeit.strip()
        json_data[counter]['location-' + lang] = location
        json_data[counter]['titel-' + lang] = titel.replace(' ', ' ').replace('L. ... o.', '').strip()
        json_data[counter]['untertitel-' + lang] = untertitel.replace(' ', ' ').replace('L. ... o.', '').strip()
        json_data[counter]['content-' + lang] = content.strip().replace(' ', ' ').replace('L. ... o.', '').strip()
        json_data[counter]['lang'] = program_lang

    block = []

    for i in range(len(paragraphs)):
        if paragraphs[i].text == '':
            check_block(block)
            block = []
            continue
        if paragraphs[i].text.startswith('***'):
            if len(block) > 0:
                check_block(block)
            break
        block.append(paragraphs[i])

    with open(json_filename, 'w') as f:
        json.dump(json_data, f, indent=4, ensure_ascii=False)
        f.close()


def combine_jsons():
    with open('static/json/cbt_programm.json') as f:
        combined = json.load(f)
        f.close()

    with open('static/json/cbt_programm_de.json') as f:
        de = json.load(f)
        f.close()
    with open('static/json/cbt_programm_hu.json') as f:
        hu = json.load(f)
        f.close()
    with open('static/json/cbt_programm_sk.json') as f:
        sk = json.load(f)
        f.close()
    with open('static/json/cbt_programm_pl.json') as f:
        pl = json.load(f)
        f.close()
    with open('static/json/cbt_programm_cs.json') as f:
        cs = json.load(f)
        f.close()

    for key in de.keys():
        #combined[key] = de[key]

        combined[key]['type'] = 0

        combined[key]['content-hu'] = hu[key]['content-hu']
        combined[key]['titel-hu'] = hu[key]['titel-hu']
        combined[key]['untertitel-hu'] = hu[key]['untertitel-hu']
        combined[key]['location-hu'] = hu[key]['location-hu']

        combined[key]['content-sk'] = sk[key]['content-sk']
        combined[key]['titel-sk'] = sk[key]['titel-sk']
        combined[key]['untertitel-sk'] = sk[key]['untertitel-sk']
        combined[key]['location-sk'] = sk[key]['location-sk']

        combined[key]['content-pl'] = pl[key]['content-pl']
        combined[key]['titel-pl'] = pl[key]['titel-pl']
        combined[key]['untertitel-pl'] = pl[key]['untertitel-pl']
        combined[key]['location-pl'] = pl[key]['location-pl']

        combined[key]['content-cs'] = cs[key]['content-cs']
        combined[key]['titel-cs'] = cs[key]['titel-cs']
        combined[key]['untertitel-cs'] = cs[key]['untertitel-cs']
        combined[key]['location-cs'] = cs[key]['location-cs']

    with open('static/json/cbt_programm.json', 'w') as f:
        json.dump(combined, f, indent=4, ensure_ascii=False)
        f.close()


def test():
    for i in range(len(paragraphs)):
        if paragraphs[i].text == '':
            if paragraphs[i + 1]:
                t = paragraphs[i + 1].text
                if t != '' and t[0] and len(t[0]) > 0 and t[0].isdigit():
                    print(t)
            if paragraphs[i + 2]:
                t = paragraphs[i + 2].runs
                for run in t:
                    if run.bold:
                        print(run.text)


#create_json2(filename_de, 'de')
#create_json2(filename_hu, 'hu')
#create_json2(filename_sk, 'sk')
#create_json2(filename_pl, 'pl')
#create_json2(filename_cs, 'cs')
combine_jsons()
